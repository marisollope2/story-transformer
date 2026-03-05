"""
File extraction utilities for various file formats
Supports .txt, .docx, and Google Docs
"""
import io
import re
from typing import Optional


def extract_text_from_file(uploaded_file) -> str:
    """
    Extract text from uploaded file based on file extension.
    
    Args:
        uploaded_file: Streamlit uploaded file object
    
    Returns:
        Extracted text as string
    """
    filename = uploaded_file.name.lower()
    
    if filename.endswith('.txt'):
        return extract_text_from_txt(uploaded_file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Supported types: .txt, .docx")


def extract_text_from_txt(uploaded_file) -> str:
    """Extract text from .txt file"""
    return uploaded_file.read().decode("utf-8")


def extract_text_from_docx(uploaded_file) -> str:
    """Extract text from .docx file"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx files. Install it with: pip install python-docx"
        )
    
    # Read the file into memory
    file_bytes = uploaded_file.read()
    doc = Document(io.BytesIO(file_bytes))
    
    # Extract text with structure
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading', '').strip()
                if level.isdigit():
                    paragraphs.append(f"Section Header: {text}")
                else:
                    paragraphs.append(f"Title: {text}")
            else:
                paragraphs.append(text)
    
    # Extract tables
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                tables_text.append(row_text)
    
    # Combine all text
    full_text = "\n\n".join(paragraphs)
    if tables_text:
        full_text += "\n\nTables:\n" + "\n".join(tables_text)
    
    return full_text


def extract_from_google_docs_url(url: str) -> Optional[str]:
    """
    Extract text from Google Docs URL.

    The document must be shared so that "Anyone with the link" can view it.
    We build a clean export URL (no extra query params) to avoid 404s from
    malformed URLs.

    Args:
        url: Google Docs URL (e.g. .../d/DOC_ID/edit?usp=sharing)

    Returns:
        Extracted text or None if extraction fails
    """
    try:
        import requests
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("requests and beautifulsoup4 are required for Google Docs extraction")

    doc_id = extract_google_doc_id(url)
    if not doc_id:
        raise ValueError("Invalid Google Docs URL format. Expected a URL like https://docs.google.com/document/d/DOCUMENT_ID/edit")

    # Build a single clean export URL. Do not append the user's query string
    # (e.g. ?usp=sharing), which can produce invalid URLs like format=txt?usp=sharing.
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"

    try:
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept": "text/plain,text/html,*/*",
        }
        response = requests.get(export_url, headers=headers, timeout=15, allow_redirects=True)
        response.raise_for_status()

        content_type = (response.headers.get("content-type") or "").lower()
        if "html" in content_type:
            soup = BeautifulSoup(response.text, "html.parser")
            return soup.get_text(separator="\n\n", strip=True)
        return response.text

    except requests.exceptions.HTTPError as e:
        if e.response is not None and e.response.status_code == 404:
            raise Exception(
                "Failed to extract from Google Docs (404). "
                "Make sure the document is shared: Share → Anyone with the link → Viewer. "
                "If it still fails, export as .docx and use File upload instead."
            ) from e
        raise Exception(
            f"Failed to extract from Google Docs. "
            f"The document may be private. Try exporting it as .docx and uploading instead. "
            f"Error: {str(e)}"
        ) from e
    except requests.exceptions.RequestException as e:
        raise Exception(
            f"Failed to extract from Google Docs. "
            f"The document may be private. Try exporting it as .docx and uploading instead. "
            f"Error: {str(e)}"
        ) from e


def extract_google_doc_id(url: str) -> Optional[str]:
    """Extract document ID from Google Docs URL"""
    # Pattern: /d/{DOC_ID}/
    match = re.search(r'/d/([a-zA-Z0-9-_]+)', url)
    if match:
        return match.group(1)
    return None


def is_google_docs_url(url: str) -> bool:
    """Check if URL is a Google Docs URL"""
    return 'docs.google.com/document' in url.lower()

